"""
Enhanced export service for multiple output formats.

This module provides comprehensive export capabilities for transcription results,
supporting various professional formats including subtitles (SRT/VTT),
documents (PDF/DOCX), and enhanced text outputs.
"""

import logging
import os
from typing import Any, Dict, List, Optional

# Optional dependencies for enhanced export formats
try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False
    print("📄 PDF export not available - install reportlab: pip install reportlab")

try:
    import docx

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("📝 DOCX export not available - install python-docx: pip install python-docx")

logger = logging.getLogger(__name__)


class EnhancedExportService:
    """
    Comprehensive export service supporting multiple professional formats.

    This service extends the basic export functionality to support industry-standard
    formats for subtitles, documents, and enhanced reports.

    Supported Formats:
    - SRT: SubRip subtitle format for video players
    - VTT: WebVTT format for web-based video players
    - PDF: Professional PDF reports with analysis
    - DOCX: Microsoft Word documents with structured content
    - Enhanced TXT: Improved text formats with better structure
    """

    def __init__(self):
        """Initialize the enhanced export service."""
        self.styles = getSampleStyleSheet() if REPORTLAB_AVAILABLE else None
        logger.info("Enhanced export service initialized")

    def export_all_formats(
        self, results: Dict[str, Any], export_options: Optional[Dict[str, bool]] = None
    ) -> Dict[str, str]:
        """
        Export transcription results in all available formats.

        Args:
            results: Complete transcription results dictionary
            export_options: Optional dictionary specifying which formats to export

        Returns:
            Dictionary mapping format names to file paths of exported files

        Example:
            >>> exporter = EnhancedExportService()
            >>> exported_files = exporter.export_all_formats(results)
            >>> print(exported_files['srt'])  # Path to SRT file
        """
        if export_options is None:
            export_options = {
                "srt": True,
                "vtt": True,
                "pdf": REPORTLAB_AVAILABLE,
                "docx": DOCX_AVAILABLE,
                "enhanced_txt": True,
            }

        session_dir = results["session_dir"]
        exported_files = {}

        # Get all segments for subtitle formats
        all_segments = self._get_all_segments(results)

        # Define all format handlers with consistent filenames
        format_handlers = {
            # Subtitle formats (require segments)
            "srt": (
                "subtitles.srt",
                lambda segments, path: self.export_to_srt(segments, path),
            ),
            "vtt": (
                "subtitles.vtt",
                lambda segments, path: self.export_to_vtt(segments, path),
            ),
            # Document formats (require full results)
            "pdf": (
                "analysis_report.pdf",
                lambda results, path: (
                    self.export_to_pdf(results, path) if REPORTLAB_AVAILABLE else None
                ),
            ),
            "docx": (
                "transcript_report.docx",
                lambda results, path: (
                    self.export_to_docx(results, path) if DOCX_AVAILABLE else None
                ),
            ),
            # Text formats (require full results)
            "enhanced_txt": (
                "transcript_enhanced.txt",
                lambda results, path: self.export_enhanced_text(results, path),
            ),
            "basic_txt": (
                "transcript.txt",
                lambda results, path: self.export_basic_text(results, path),
            ),
            # Data formats (require full results)
            "json": (
                "analysis.json",
                lambda results, path: self.export_to_json(results, path),
            ),
            "html": (
                "searchable_transcript.html",
                lambda results, path: self.export_to_html(results, path),
            ),
        }

        # Single consolidated export loop
        for format_name, (file_name, handler) in format_handlers.items():
            if export_options.get(format_name, True) and handler:
                try:
                    file_path = os.path.join(session_dir, file_name)

                    # Determine the correct input based on format type
                    if format_name in ["srt", "vtt"]:
                        # Subtitle formats need segments
                        handler(all_segments, file_path)
                    else:
                        # All other formats need full results
                        handler(results, file_path)

                    exported_files[format_name] = file_path
                    logger.info(f"Exported {format_name.upper()} to {file_path}")
                except Exception as e:
                    logger.error(f"Failed to export {format_name.upper()}: {e}")

        return exported_files

    def export_to_srt(self, segments: List[Dict[str, Any]], output_path: str) -> None:
        """
        Export segments as SRT subtitle file.

        Args:
            segments: List of timestamped segments
            output_path: Path where SRT file will be saved

        SRT Format:
            1
            00:00:00,000 --> 00:00:03,000
            This is the first subtitle

            2
            00:00:03,000 --> 00:00:06,000
            This is the second subtitle
        """
        with open(output_path, "w", encoding="utf-8") as f:
            for i, segment in enumerate(segments, 1):
                start_time = self._format_srt_time(segment["start"])
                end_time = self._format_srt_time(segment["end"])

                f.write(f"{i}\n")
                f.write(f"{start_time} --> {end_time}\n")
                f.write(f"{segment['text'].strip()}\n\n")

        logger.info(f"Exported {len(segments)} segments to SRT: {output_path}")

    def export_to_vtt(self, segments: List[Dict[str, Any]], output_path: str) -> None:
        """
        Export segments as WebVTT subtitle file.

        Args:
            segments: List of timestamped segments
            output_path: Path where VTT file will be saved

        VTT Format:
            WEBVTT

            00:00:00.000 --> 00:00:03.000
            This is the first subtitle

            00:00:03.000 --> 00:00:06.000
            This is the second subtitle
        """
        with open(output_path, "w", encoding="utf-8") as f:
            f.write("WEBVTT\n\n")

            for segment in segments:
                start_time = self._format_vtt_time(segment["start"])
                end_time = self._format_vtt_time(segment["end"])

                f.write(f"{start_time} --> {end_time}\n")
                f.write(f"{segment['text'].strip()}\n\n")

        logger.info(f"Exported {len(segments)} segments to VTT: {output_path}")

    def export_to_pdf(self, results: Dict[str, Any], output_path: str) -> None:
        """
        Export comprehensive analysis as PDF document.

        Args:
            results: Complete transcription results
            output_path: Path where PDF will be saved

        Features:
        - Professional document layout
        - Summary statistics
        - Analysis highlights
        - Full transcript with formatting
        """
        if not REPORTLAB_AVAILABLE:
            raise ImportError("ReportLab not available - cannot export PDF")

        doc = SimpleDocTemplate(output_path, pagesize=A4)
        story = []

        # Title page
        title = Paragraph("Video Transcription Analysis Report", self.styles["Title"])
        story.append(title)
        story.append(Spacer(1, 12))

        # Metadata section
        metadata = results.get("metadata", {})
        story.append(
            Paragraph(
                f"<b>Video:</b> {metadata.get('original_filename', 'Unknown')}",
                self.styles["Heading2"],
            )
        )

        # Format duration if available
        duration_text = "Unknown"
        if "processing_time" in metadata:
            duration_minutes = metadata["processing_time"] / 60
            duration_text = f"{duration_minutes:.1f} minutes"

        story.append(
            Paragraph(f"<b>Duration:</b> {duration_text}", self.styles["Normal"])
        )
        story.append(
            Paragraph(
                f"<b>Processed:</b> {metadata.get('created_at', 'Unknown')}",
                self.styles["Normal"],
            )
        )
        story.append(Spacer(1, 12))

        # Summary statistics
        analysis = results.get("analysis", {})
        story.append(Paragraph("Summary Statistics", self.styles["Heading2"]))
        story.append(
            Paragraph(
                f"<b>Total Words:</b> {analysis.get('total_words', 0)}",
                self.styles["Normal"],
            )
        )
        story.append(
            Paragraph(
                f"<b>Keywords Found:</b> {len(analysis.get('keyword_matches', []))}",
                self.styles["Normal"],
            )
        )
        story.append(
            Paragraph(
                f"<b>Questions Detected:</b> {len(analysis.get('questions', []))}",
                self.styles["Normal"],
            )
        )
        story.append(
            Paragraph(
                f"<b>Emphasis Cues:</b> {len(analysis.get('emphasis_cues', []))}",
                self.styles["Normal"],
            )
        )
        story.append(Spacer(1, 12))

        # Analysis highlights
        if analysis.get("keyword_matches"):
            story.append(Paragraph("Key Findings", self.styles["Heading2"]))
            for match in analysis["keyword_matches"][:5]:  # Top 5 keywords
                story.append(
                    Paragraph(
                        f"<b>{match['keyword']}</b>: {match['count']} mentions",
                        self.styles["Normal"],
                    )
                )
            story.append(Spacer(1, 12))

        # Questions section
        if analysis.get("questions"):
            story.append(Paragraph("Questions Detected", self.styles["Heading2"]))
            for question in analysis["questions"][:10]:  # First 10 questions
                story.append(
                    Paragraph(
                        f"[{question['timestamp']}] {question['text']}",
                        self.styles["Normal"],
                    )
                )
            story.append(Spacer(1, 12))

        # Full transcript (truncated for PDF)
        story.append(Paragraph("Transcript Excerpt", self.styles["Heading2"]))
        transcript_text = results.get("full_transcript", "")
        if len(transcript_text) > 2000:
            transcript_text = transcript_text[:2000] + "... (truncated)"

        story.append(Paragraph(transcript_text, self.styles["Normal"]))

        # Build PDF
        doc.build(story)
        logger.info(f"Exported PDF report: {output_path}")

    def export_to_docx(self, results: Dict[str, Any], output_path: str) -> None:
        """
        Export as Microsoft Word document.

        Args:
            results: Complete transcription results
            output_path: Path where DOCX will be saved

        Features:
        - Structured document with headings
        - Summary statistics table
        - Timestamped transcript
        - Analysis sections
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx not available - cannot export DOCX")

        doc = docx.Document()

        # Title and metadata
        doc.add_heading("Video Transcription Analysis", 0)

        metadata = results.get("metadata", {})
        doc.add_heading("Video Information", level=1)

        info_table = doc.add_table(rows=3, cols=2)
        info_table.style = "Table Grid"

        # Add metadata to table
        info_table.cell(0, 0).text = "Filename"
        info_table.cell(0, 1).text = metadata.get("original_filename", "Unknown")
        info_table.cell(1, 0).text = "Session ID"
        info_table.cell(1, 1).text = metadata.get("session_id", "Unknown")
        info_table.cell(2, 0).text = "Processed"
        info_table.cell(2, 1).text = metadata.get("created_at", "Unknown")

        # Analysis summary
        analysis = results.get("analysis", {})
        doc.add_heading("Analysis Summary", level=1)

        summary_table = doc.add_table(rows=4, cols=2)
        summary_table.style = "Table Grid"
        summary_table.cell(0, 0).text = "Total Words"
        summary_table.cell(0, 1).text = str(analysis.get("total_words", 0))
        summary_table.cell(1, 0).text = "Keywords Found"
        summary_table.cell(1, 1).text = str(len(analysis.get("keyword_matches", [])))
        summary_table.cell(2, 0).text = "Questions Detected"
        summary_table.cell(2, 1).text = str(len(analysis.get("questions", [])))
        summary_table.cell(3, 0).text = "Emphasis Cues"
        summary_table.cell(3, 1).text = str(len(analysis.get("emphasis_cues", [])))

        # Keywords section
        if analysis.get("keyword_matches"):
            doc.add_heading("Keywords Found", level=1)
            for match in analysis["keyword_matches"]:
                keyword_para = doc.add_paragraph()
                keyword_para.add_run(f"{match['keyword']}: ").bold = True
                keyword_para.add_run(f"{match['count']} mentions")

        # Questions section
        if analysis.get("questions"):
            doc.add_heading("Questions Detected", level=1)
            for question in analysis["questions"]:
                question_para = doc.add_paragraph()
                question_para.add_run(f"[{question['timestamp']}] ").bold = True
                question_para.add_run(question["text"])

        # Transcript section
        doc.add_heading("Full Transcript", level=1)

        # Get all segments for timestamped transcript
        all_segments = self._get_all_segments(results)

        for segment in all_segments[
            :50
        ]:  # Limit to first 50 segments for document size
            transcript_para = doc.add_paragraph()
            timestamp_run = transcript_para.add_run(f"[{segment['timestamp_str']}] ")
            timestamp_run.bold = True
            timestamp_run.font.color.rgb = docx.shared.RGBColor(102, 126, 234)
            transcript_para.add_run(segment["text"])

        if len(all_segments) > 50:
            doc.add_paragraph("... (transcript truncated for document size)")

        # Save document
        doc.save(output_path)
        logger.info(f"Exported DOCX document: {output_path}")

    def export_enhanced_text(self, results: Dict[str, Any], output_path: str) -> None:
        """
        Export enhanced text format with better structure and formatting.

        Args:
            results: Complete transcription results
            output_path: Path where enhanced text will be saved

        Features:
        - Clear section headers
        - Formatted timestamps
        - Analysis highlights
        - Better readability
        """
        with open(output_path, "w", encoding="utf-8") as f:
            # Header
            f.write("=" * 80 + "\n")
            f.write("VIDEO TRANSCRIPTION ANALYSIS REPORT\n")
            f.write("=" * 80 + "\n\n")

            # Metadata
            metadata = results.get("metadata", {})
            f.write("📋 SESSION INFORMATION\n")
            f.write("-" * 40 + "\n")
            f.write(f"Video File: {metadata.get('original_filename', 'Unknown')}\n")
            f.write(f"Session ID: {metadata.get('session_id', 'Unknown')}\n")
            f.write(f"Processed: {metadata.get('created_at', 'Unknown')}\n")
            if "processing_time" in metadata:
                f.write(f"Processing Time: {metadata['processing_time']:.1f} seconds\n")
            f.write("\n")

            # Summary statistics
            analysis = results.get("analysis", {})
            f.write("📊 SUMMARY STATISTICS\n")
            f.write("-" * 40 + "\n")
            f.write(f"Total Words: {analysis.get('total_words', 0):,}\n")
            f.write(f"Keywords Found: {len(analysis.get('keyword_matches', []))}\n")
            f.write(f"Questions Detected: {len(analysis.get('questions', []))}\n")
            f.write(f"Emphasis Cues: {len(analysis.get('emphasis_cues', []))}\n")
            f.write("\n")

            # Keywords section
            if analysis.get("keyword_matches"):
                f.write("🔍 KEYWORDS ANALYSIS\n")
                f.write("-" * 40 + "\n")
                for match in analysis["keyword_matches"]:
                    f.write(f"• {match['keyword']}: {match['count']} mentions\n")
                f.write("\n")

            # Questions section
            if analysis.get("questions"):
                f.write("❓ QUESTIONS DETECTED\n")
                f.write("-" * 40 + "\n")
                for question in analysis["questions"]:
                    f.write(f"[{question['timestamp']}] {question['text']}\n")
                f.write("\n")

            # Emphasis cues section
            if analysis.get("emphasis_cues"):
                f.write("⚡ EMPHASIS CUES\n")
                f.write("-" * 40 + "\n")
                for cue in analysis["emphasis_cues"]:
                    f.write(f"[{cue['timestamp']}] {cue['text']}\n")
                f.write("\n")

            # Full transcript
            f.write("📝 FULL TRANSCRIPT\n")
            f.write("-" * 40 + "\n")

            all_segments = self._get_all_segments(results)
            for segment in all_segments:
                f.write(f"[{segment['timestamp_str']}] {segment['text']}\n")

        logger.info(f"Exported enhanced text: {output_path}")

    def export_basic_text(self, results: Dict[str, Any], output_path: str) -> None:
        """
        Export basic plain text transcript.

        Args:
            results: Transcription results
            output_path: Path where text file will be saved
        """
        with open(output_path, "w", encoding="utf-8") as f:
            # Simple transcript
            all_segments = self._get_all_segments(results)
            for segment in all_segments:
                f.write(f"{segment['text']}\n")

        logger.info(f"Exported basic text: {output_path}")

    def export_to_json(self, results: Dict[str, Any], output_path: str) -> None:
        """
        Export complete analysis results to JSON format.

        Args:
            results: Complete transcription and analysis results
            output_path: Path where JSON file should be saved
        """
        try:
            import json

            with open(output_path, "w", encoding="utf-8") as f:
                json.dump(results, f, indent=2, ensure_ascii=False)
            logger.info(f"JSON export completed: {output_path}")
        except Exception as e:
            logger.error(f"Failed to export JSON: {e}")
            raise

    def export_to_html(self, results: Dict[str, Any], output_path: str) -> None:
        """
        Export interactive searchable transcript to HTML format.

        Args:
            results: Complete transcription and analysis results
            output_path: Path where HTML file should be saved
        """
        try:
            # Extract transcript segments
            segments = results.get("segments", [])
            metadata = results.get("metadata", {})
            filename = metadata.get("filename", "Unknown")

            html_content = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Searchable Transcript - {filename}</title>
    <style>
        body {{
            font-family: Arial, sans-serif;
            max-width: 800px;
            margin: 0 auto;
            padding: 20px;
            line-height: 1.6;
        }}
        .header {{
            border-bottom: 2px solid #333;
            padding-bottom: 10px;
            margin-bottom: 20px;
        }}
        .search-box {{
            width: 100%;
            padding: 10px;
            margin-bottom: 20px;
            border: 2px solid #ddd;
            border-radius: 5px;
            font-size: 16px;
        }}
        .segment {{
            margin-bottom: 15px;
            padding: 10px;
            border-left: 3px solid #007acc;
            background-color: #f9f9f9;
        }}
        .timestamp {{
            font-weight: bold;
            color: #007acc;
            margin-bottom: 5px;
        }}
        .text {{
            margin-bottom: 5px;
        }}
        .highlight {{
            background-color: yellow;
        }}
        .metadata {{
            background-color: #f0f0f0;
            padding: 15px;
            border-radius: 5px;
            margin-bottom: 20px;
        }}
    </style>
</head>
<body>
    <div class="header">
        <h1>Searchable Transcript</h1>
        <h2>{filename}</h2>
    </div>

    <div class="metadata">
        <h3>File Information</h3>
        <p><strong>Duration:</strong> {metadata.get('duration', 'Unknown')}</p>
        <p><strong>Created:</strong> {metadata.get('created_date', 'Unknown')}</p>
    </div>

    <input type="text" class="search-box" id="searchInput" placeholder="Search transcript...">

    <div id="transcript">
"""

            # Add transcript segments
            for i, segment in enumerate(segments):
                start_time = self._format_timestamp(segment.get("start", 0))
                end_time = self._format_timestamp(segment.get("end", 0))
                text = segment.get("text", "").strip()

                html_content += f"""
        <div class="segment" data-text="{text.lower()}">
            <div class="timestamp">[{start_time} - {end_time}]</div>
            <div class="text">{text}</div>
        </div>
"""

            # Add JavaScript for search functionality
            html_content += """
    </div>

    <script>
        document.getElementById('searchInput').addEventListener('input', function(e) {
            const searchTerm = e.target.value.toLowerCase();
            const segments = document.querySelectorAll('.segment');

            segments.forEach(segment => {
                const text = segment.getAttribute('data-text');
                const textDiv = segment.querySelector('.text');

                if (searchTerm === '' || text.includes(searchTerm)) {
                    segment.style.display = 'block';

                    // Highlight search term
                    if (searchTerm !== '') {
                        const originalText = textDiv.textContent;
                        const highlightedText = originalText.replace(
                            new RegExp(searchTerm, 'gi'),
                            '<span class="highlight">$&</span>'
                        );
                        textDiv.innerHTML = highlightedText;
                    } else {
                        textDiv.innerHTML = textDiv.textContent;
                    }
                } else {
                    segment.style.display = 'none';
                }
            });
        });
    </script>
</body>
</html>
"""

            with open(output_path, "w", encoding="utf-8") as f:
                f.write(html_content)
            logger.info(f"HTML export completed: {output_path}")
        except Exception as e:
            logger.error(f"Failed to export HTML: {e}")
            raise

    def _format_timestamp(self, seconds: float) -> str:
        """
        Format seconds as HH:MM:SS timestamp.

        Args:
            seconds: Time in seconds

        Returns:
            Formatted timestamp string
        """
        hours = int(seconds // 3600)
        minutes = int((seconds % 3600) // 60)
        secs = int(seconds % 60)
        return f"{hours:02d}:{minutes:02d}:{secs:02d}"

    def _get_all_segments(self, results: Dict[str, Any]) -> List[Dict[str, Any]]:
        """Extract and sort all segments from results."""
        all_segments = []

        # Handle both chunk-based and direct segments
        if "chunks" in results:
            for chunk in results["chunks"]:
                if isinstance(chunk, list):
                    all_segments.extend(chunk)
                elif isinstance(chunk, dict) and "segments" in chunk:
                    all_segments.extend(chunk["segments"])
        elif "segments" in results:
            all_segments = results["segments"]

        # Sort by timestamp
        all_segments.sort(key=lambda x: x.get("start", 0))
        return all_segments

    def _format_srt_time(self, seconds: float) -> str:
        """Format time for SRT format (HH:MM:SS,mmm)."""
        hours = int(seconds // 3600)
        minutes = int((seconds % 3600) // 60)
        secs = int(seconds % 60)
        millisecs = int((seconds % 1) * 1000)
        return f"{hours:02d}:{minutes:02d}:{secs:02d},{millisecs:03d}"

    def _format_vtt_time(self, seconds: float) -> str:
        """Format time for VTT format (HH:MM:SS.mmm)."""
        hours = int(seconds // 3600)
        minutes = int((seconds % 3600) // 60)
        secs = int(seconds % 60)
        millisecs = int((seconds % 1) * 1000)
        return f"{hours:02d}:{minutes:02d}:{secs:02d}.{millisecs:03d}"

    def get_available_formats(self) -> Dict[str, bool]:
        """
        Get dictionary of available export formats.

        Returns:
            Dictionary mapping format names to availability status
        """
        return {
            "srt": True,
            "vtt": True,
            "pdf": REPORTLAB_AVAILABLE,
            "docx": DOCX_AVAILABLE,
            "enhanced_txt": True,
            "basic_txt": True,
            "json": True,
            "html": True,
        }

    def get_format_descriptions(self) -> Dict[str, str]:
        """
        Get descriptions of available export formats.

        Returns:
            Dictionary mapping format names to descriptions
        """
        return {
            "srt": "SubRip subtitle format for video players",
            "vtt": "WebVTT format for web-based video players",
            "pdf": "Professional PDF report with analysis (requires reportlab)",
            "docx": "Microsoft Word document (requires python-docx)",
            "enhanced_txt": "Structured text format with improved readability",
            "basic_txt": "Simple plain text transcript",
            "json": "Complete analysis data in JSON format",
            "html": "Interactive searchable transcript",
        }
